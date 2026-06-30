"""Build PROJECT_PROPOSAL_REPORT.docx with cover page, TOC, ERD image, and page numbers."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

DOCS = Path(__file__).resolve().parent
OUTPUT = DOCS / "PROJECT_PROPOSAL_REPORT.docx"
ERD_IMAGE = DOCS / "esukan-erd.png"


def set_document_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.color.rgb = RGBColor(0x28, 0x48, 0x78)


def add_page_number_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def add_cover_page(doc: Document) -> None:
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("E-SUKAN")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x28, 0x48, 0x78)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub.add_run("Project Proposal Report")
    s.font.size = Pt(18)
    s.bold = True

    doc.add_paragraph()
    lines = [
        "CSC584 — Group Project",
        "Campus Facility & Equipment Booking System",
        "",
        "Information System Development",
        "",
        "Group Members:",
        "[Name 1] — [Matric No.]",
        "[Name 2] — [Matric No.]",
        "[Name 3] — [Matric No.]",
        "[Name 4] — [Matric No.]",
        "",
        f"Submission Date: {date.today().strftime('%d %B %Y')}",
        "Institution: [University / Faculty Name]",
        "Lecturer: [Lecturer Name]",
    ]
    for line in lines:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if line.startswith("["):
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()


def add_table_of_contents(doc: Document) -> None:
    doc.add_heading("Table of Contents", level=1)
    entries = [
        ("1.0", "Project Overview"),
        ("2.0", "Problem Statement"),
        ("3.0", "Objective"),
        ("4.0", "Proposed User Interface"),
        ("5.0", "Database Design"),
        ("6.0", "Flow of Application"),
        ("", "Appendix A — Scoresheet Mapping"),
        ("", "Appendix B — Project Files"),
    ]
    for num, title in entries:
        p = doc.add_paragraph()
        if num:
            p.add_run(f"{num}\t{title}")
        else:
            p.add_run(title)
    doc.add_paragraph(
        "Tip: In Microsoft Word, select all and press F9 to refresh page numbers "
        "after adding screenshots."
    ).italic = True
    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
    doc.add_paragraph()


def add_screenshot_placeholder(doc: Document, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ INSERT SCREENSHOT: {caption} ]")
    run.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x33, 0x33)
    doc.add_paragraph()


def add_section_1(doc: Document) -> None:
    add_heading(doc, "1.0 PROJECT OVERVIEW", 1)
    doc.add_paragraph(
        "E-Sukan is a web-based information system for managing campus sports facilities "
        "and equipment at a university. Students and lecturers book sports courts "
        "(badminton and futsal), rent equipment, join waitlists when slots are full, "
        "pay for bookings and rentals, and participate in tournaments. Administrators "
        "manage users, facilities, equipment, bookings, and system settings."
    )
    add_heading(doc, "Architecture", 2)
    add_table(
        doc,
        ["Layer", "Description"],
        [
            ["Presentation", "Responsive dashboard (login.html, index.html), custom CSS, sidebar navigation, REST API calls"],
            ["Business logic", "Jakarta Servlet 6 JSON APIs (/api/bookings, /api/equipment, /api/tournaments, etc.)"],
            ["Data", "MySQL/H2 database esukan_db, 14 tables, HikariCP pool, sql/schema.sql"],
        ],
    )
    add_heading(doc, "User roles", 2)
    add_table(
        doc,
        ["Role", "Capabilities"],
        [
            ["STUDENT", "Book facilities, rent equipment, waitlist, payments, tournament registration, dashboard reports"],
            ["LECTURER", "Student features plus create and manage tournaments"],
            ["ADMIN", "User management, facility/equipment CRUD, booking approval, system settings"],
        ],
    )
    doc.add_paragraph(
        "Authentication uses JWT login, BCrypt passwords, password reset tokens, and "
        "role-based menu visibility on every screen."
    )


def add_section_2(doc: Document) -> None:
    add_heading(doc, "2.0 PROBLEM STATEMENT", 1)
    problems = [
        "Double booking and scheduling conflicts — no real-time availability check.",
        "Poor equipment tracking — quantity, condition, and return status not centralized.",
        "No structured waitlist when facility slots are full.",
        "Weak payment and cost visibility — deposits and booking fees hard to reconcile.",
        "Tournament management disconnected from facility and user data.",
        "Limited reporting — no dashboard for peak hours, inventory health, or trends.",
        "Security gaps — informal sign-in without proper session control and validation.",
    ]
    add_bullets(doc, problems)
    doc.add_paragraph(
        "These problems reduce efficiency, increase administrative workload, and "
        "create a poor experience for campus sports users."
    )


def add_section_3(doc: Document) -> None:
    add_heading(doc, "3.0 OBJECTIVE", 1)
    objectives = [
        "Provide a centralized online platform for facility booking with conflict prevention.",
        "Implement equipment rental management with quantity, status, deposit, and returns.",
        "Support a booking waitlist with promotion when bookings are cancelled.",
        "Record payments linked to bookings and rentals (pending, paid, failed).",
        "Enable tournament lifecycle management integrated with users and facilities.",
        "Deliver analysis reports: booking stats, inventory health, peak usage hours.",
        "Enforce secure access via JWT, roles, and server-side input validation.",
        "Offer a consistent responsive UI with navigation on every screen.",
    ]
    for i, obj in enumerate(objectives, 1):
        doc.add_paragraph(f"{i}. {obj}")


def add_section_4(doc: Document) -> None:
    add_heading(doc, "4.0 PROPOSED USER INTERFACE", 1)
    doc.add_paragraph(
        "The UI uses a custom responsive template (CSS variables, sidebar layout, cards) "
        "in src/main/resources/static/. Capture screenshots from http://localhost:9090 "
        "after running mvn jetty:run."
    )
    screens = [
        ("UI 1 — Login & Registration (login.html)", "Login page — sign in, register, forgot password"),
        ("UI 2 — Dashboard", "Dashboard — stats, inventory health report, peak hours, schedule"),
        ("UI 3 — Facilities & Booking", "Facilities list and New Booking modal with payment"),
        ("UI 4 — Equipment & Rentals", "Equipment catalog and active/returned rentals"),
        ("UI 5 — Tournaments", "Tournament list, team registration, match bracket"),
        ("UI 6 — User Management (Admin)", "Admin user table — create, edit role, enable/disable"),
        ("UI 7 — Password Reset (reset.html)", "Reset password form with token"),
    ]
    for title, caption in screens:
        add_heading(doc, title, 2)
        add_screenshot_placeholder(doc, caption)
    add_heading(doc, "Navigation consistency", 2)
    add_bullets(
        doc,
        [
            "Fixed sidebar with role-filtered menu on every authenticated page.",
            "Top bar with page title and primary action (+ New Booking, etc.).",
            "User chip and logout in sidebar footer.",
        ],
    )


def add_section_5(doc: Document) -> None:
    add_heading(doc, "5.0 DATABASE DESIGN", 1)
    doc.add_paragraph(
        "Database name: esukan_db (14 tables). Source: sql/schema.sql. "
        "ERD tool: dbdiagram.io using docs/esukan-erd.dbml."
    )
    add_heading(doc, "5.1 Table structure", 2)
    add_table(
        doc,
        ["Table", "Purpose", "Key attributes"],
        [
            ["system_settings", "Config", "setting_key PK, setting_value"],
            ["users", "Accounts", "id PK, username UK, email UK, role, password_hash"],
            ["password_reset_tokens", "Recovery", "token UK, user_id FK"],
            ["facilities", "Courts", "name, type, open_time, close_time, cost_per_hour"],
            ["equipment", "Inventory", "name, category, status, quantity"],
            ["facility_equipment", "M:N link", "facility_id, equipment_id composite PK"],
            ["bookings", "Reservations", "facility_id FK, user_id FK, booking_date, status"],
            ["booking_waitlist", "FIFO queue", "facility_id FK, promoted_booking_id"],
            ["equipment_rentals", "Checkout", "equipment_id FK, deposit_amount, status"],
            ["payments", "Transactions", "booking_id FK, rental_id FK, amount, status"],
            ["tournaments", "Events", "organizer_id FK, venue_facility_id FK, format"],
            ["tournament_registrations", "Teams", "tournament_id FK, team_name"],
            ["tournament_team_members", "Roster", "registration_id FK"],
            ["tournament_matches", "Bracket", "team_a/b/winner FK, next_match_id self-FK"],
        ],
    )
    add_heading(doc, "5.2 Entity Relationship Diagram", 2)
    if ERD_IMAGE.exists():
        doc.add_paragraph("Figure 1: E-Sukan database ERD (orthogonal notation: 1 = one, crow's foot = many).")
        doc.add_picture(str(ERD_IMAGE), width=Inches(6.8))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        add_screenshot_placeholder(doc, "ERD diagram (docs/esukan-erd.png)")
    add_heading(doc, "5.3 Relationships between tables", 2)
    add_table(
        doc,
        ["Parent (1)", "Child (many)", "Description"],
        [
            ["users", "bookings, rentals, tokens, tournaments", "One user, many records"],
            ["facilities", "bookings, waitlist, tournaments", "One facility, many uses"],
            ["equipment", "rentals, facility_equipment", "One item, many rentals"],
            ["bookings / rentals", "payments", "Optional payment link"],
            ["tournaments", "registrations, matches", "One event, many teams/matches"],
            ["tournament_matches", "tournament_matches", "Bracket next_match_id self-reference"],
        ],
    )


def add_section_6(doc: Document) -> None:
    add_heading(doc, "6.0 FLOW OF APPLICATION", 1)

    add_heading(doc, "6.1 High-level system flow", 2)
    add_bullets(
        doc,
        [
            "User opens login.html → Sign in or Register → JWT stored in browser.",
            "Authenticated user opens index.html Dashboard.",
            "Navigate via sidebar: Facilities, Bookings, Equipment, Rentals, Tournaments, Users (admin).",
            "Dashboard loads analysis reports (stats, health, peak hours).",
        ],
    )

    add_heading(doc, "6.2 Create / Read (CR)", 2)
    add_table(
        doc,
        ["Module", "Create", "Read"],
        [
            ["Auth", "POST /api/auth/register", "POST login, GET /api/auth/me"],
            ["Facilities", "Admin POST facility", "GET /api/facilities"],
            ["Bookings", "POST /api/bookings + payment", "GET bookings, GET dashboard"],
            ["Equipment", "Admin POST equipment", "GET catalog, health report"],
            ["Rentals", "POST /api/rentals", "GET active/returned rentals"],
            ["Tournaments", "POST tournament, register team", "GET tournaments, matches"],
            ["Users", "Admin POST user", "GET user list"],
        ],
    )
    doc.add_paragraph("Success message shown in UI after each successful create operation.")

    add_heading(doc, "6.3 Update", 2)
    add_bullets(
        doc,
        [
            "Admin edits facility/equipment → PUT/PATCH API → success message.",
            "Admin approves booking → status CONFIRMED → confirmation dialog first.",
            "Lecturer changes tournament status (DRAFT → OPEN → CLOSED).",
            "User returns equipment → rental status RETURNED.",
        ],
    )

    add_heading(doc, "6.4 Delete", 2)
    add_bullets(
        doc,
        [
            "User/admin selects delete → confirmation dialog → DELETE API → refresh list.",
            "Applies to: booking cancel, facility/equipment (admin), user (admin), tournament.",
        ],
    )

    add_heading(doc, "6.5 Report (Analysis)", 2)
    add_bullets(
        doc,
        [
            "Open Dashboard → GET /api/bookings/dashboard (total, today, pending).",
            "GET /api/equipment/health-report (available, damaged, maintenance).",
            "Peak usage hours chart per selected facility.",
            "Today's booking schedule table.",
        ],
    )


def add_appendices(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "APPENDIX A — Scoresheet Mapping", 1)
    add_table(
        doc,
        ["Scoresheet", "Criterion", "Max", "Section"],
        [
            ["Proposal", "Overview, problem, objective", "5", "1.0 – 3.0"],
            ["Proposal", "Design / UI template", "5", "4.0"],
            ["Proposal", "Table structure", "3", "5.1"],
            ["Proposal", "Relationships", "3", "5.2 – 5.3"],
            ["Proposal", "CRUD + Report flows", "4", "6.2 – 6.5"],
            ["Information System", "Validation + session", "4", "JWT, servlets"],
            ["Information System", "CRUD messages + reports", "8", "app.js + dashboard"],
            ["Information System", "UI + navigation", "8", "style.css sidebar"],
            ["Presentation", "Video", "10", "Demo all modules"],
        ],
    )
    add_heading(doc, "APPENDIX B — Project Files", 1)
    add_table(
        doc,
        ["Item", "Path"],
        [
            ["ERD image", "docs/esukan-erd.png"],
            ["DBML (dbdiagram.io)", "docs/esukan-erd.dbml"],
            ["SQL schema", "sql/schema.sql"],
            ["This report (Word)", "docs/PROJECT_PROPOSAL_REPORT.docx"],
            ["README / run guide", "README.md"],
        ],
    )


def configure_footer(doc: Document) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    first_footer = section.first_page_footer
    first_footer.paragraphs[0].text = ""

    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ""
    run = p.add_run("E-Sukan — Project Proposal    |    Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    add_page_number_field(p)


def build() -> Path:
    doc = Document()
    set_document_defaults(doc)
    configure_footer(doc)
    add_cover_page(doc)
    add_table_of_contents(doc)
    add_section_1(doc)
    add_section_2(doc)
    add_section_3(doc)
    add_section_4(doc)
    doc.add_page_break()
    add_section_5(doc)
    doc.add_page_break()
    add_section_6(doc)
    add_appendices(doc)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build()
    print(f"Created: {path}")
